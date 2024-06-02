from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pydantic import BaseModel
from typing import List

class Item(BaseModel):
    episode: int
    start: int
    end: int
    duration: int
    type: str

def generate_doc(sec: int, items: List[Item], word_name: str) -> str:
    doc = Document()
    
    # Добавляем заголовок документа
    title = doc.add_paragraph()
    title_run = title.add_run('Отчет для врача по полисомнографической записи')
    title_run.font.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Добавляем информацию о длительности анализа
    duration = doc.add_paragraph()
    duration_run = duration.add_run(f'Проведен анализ ПСГ длительностью {round(sec / 60)} минут')
    duration_run.font.size = Pt(12)
    duration.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Добавляем информацию о канале анализа
    channel = doc.add_paragraph()
    channel_run = channel.add_run('Анализировался канал "Fp1-M2"')
    channel_run.font.size = Pt(12)
    channel.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Добавляем пустую строку после информации о канале
    doc.add_paragraph()

    # Добавляем заголовок таблицы
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'  # Применяем стиль с сеткой

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер аномалии'
    hdr_cells[1].text = 'Начало(с)'
    hdr_cells[2].text = 'Конец(с)'
    hdr_cells[3].text = 'Продолжительность(с)'
    hdr_cells[4].text = 'Тип'

    # Стилизуем заголовок
    for cell in hdr_cells:
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Добавляем данные в таблицу
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.episode)
        row_cells[1].text = str(item.start)
        row_cells[2].text = str(item.end)
        row_cells[3].text = str(item.duration)
        row_cells[4].text = str(item.type)

    # Сохраняем документ
    doc.save(word_name)
    return word_name

#generate_doc(3600, [Item(episode=1, start=2, end=3, duration=4, type="test"), Item(episode=2, start=2, end=3, duration=4, type="test")], "test.docx")